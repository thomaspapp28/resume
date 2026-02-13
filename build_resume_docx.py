"""
Build a resume Word document from context.txt.
Run: python build_resume_docx.py
Output: resume.docx (in the same folder)
"""

import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import parse_xml
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call(["pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import parse_xml

FONT_NAME = "Times New Roman"
# Content (body, bullets, skills, job title, contact) font size in points
CONTENT_FONT_SIZE = 9.5
# Accent color for name, title, section titles, job titles (#0070C0)
ACCENT_COLOR = RGBColor(0x00, 0x70, 0xC0)
# Spacing between sections (before section header)
SECTION_SPACE_BEFORE = Pt(12)
# Space after section heading (title)
HEADING_SPACE_AFTER = Pt(6)
# Space between profile summary and skills block
PROFILE_SKILLS_SPACE = Pt(8)
# Skills section: hanging indent (first line at margin, wrap at 0.5 in)
SKILL_HANGING_INDENT = Inches(0.5)
# Space between each company block in work experience
COMPANY_BLOCK_SPACE = Pt(10)
# Space between subtitle (title) and contact line
TITLE_CONTACT_SPACE = Pt(6)
# Content width for right tab (A4 minus left/right margins 0.6")
LINE_WIDTH = Inches(8.27 - 0.6 - 0.6)


def load_context(path: str) -> str:
    """Load resume text from context.txt."""
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def add_heading_style(doc, name, font_size_pt=14, bold=True):
    """Add or get a heading style."""
    try:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        style = doc.styles[name]
    style.font.size = Pt(font_size_pt)
    style.font.bold = bold
    return style


def add_border_below(paragraph):
    """Add a line border below the paragraph (section header line). sz=12 → 1.5pt (a little thick)."""
    pPr = paragraph._element.get_or_add_pPr()
    bdr_xml = (
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:bottom w:val="single" w:sz="12" w:space="1" w:color="000000"/>'
        "</w:pBdr>"
    )
    pBdr = parse_xml(bdr_xml)
    pPr.append(pBdr)


def split_date_location(line):
    """Split 'May 2025 – Present    USA - Remote' into (period, location)."""
    parts = re.split(r"[\t ]{2,}", line.strip(), maxsplit=1)
    period = parts[0].strip() if parts else ""
    location = parts[1].strip() if len(parts) > 1 else ""
    return period, location


def split_contact(line):
    """Split contact line into (email, phone, location)."""
    parts = re.split(r"[\t ]{2,}", line.strip())
    parts = [p.strip() for p in parts if p.strip()]
    email = parts[0] if len(parts) > 0 else ""
    phone = parts[1] if len(parts) > 1 else ""
    location = parts[2] if len(parts) > 2 else ""
    return email, phone, location


def split_cert_line(line):
    """Split certification line into (cert_name, issuer). Handles leading '-' bullet and tab/single tab."""
    content = line.strip().lstrip("-• \t")
    parts = re.split(r"\t|[\t ]{2,}", content, maxsplit=1)
    cert_name = parts[0].strip() if parts else ""
    issuer = parts[1].strip() if len(parts) > 1 else ""
    return cert_name, issuer


def build_resume_docx(context_path: str, output_path: str) -> None:
    """Create a Word resume from context.txt content."""
    text = load_context(context_path)
    lines = [line for line in text.split("\n")]

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(CONTENT_FONT_SIZE)
    style.paragraph_format.space_after = Pt(0)

    # --- Name (title) ---
    if lines:
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(lines[0].strip())
        name_run.bold = True
        name_run.font.size = Pt(18)
        name_run.font.name = FONT_NAME
        name_run.font.color.rgb = ACCENT_COLOR
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Subtitle (title + contact) ---
    if len(lines) > 1:
        sub = doc.add_paragraph(lines[1].strip())
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in sub.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(CONTENT_FONT_SIZE)
            run.font.color.rgb = ACCENT_COLOR
    if len(lines) > 3 and lines[3].strip():
        email, phone, location = split_contact(lines[3])
        contact = doc.add_paragraph()
        contact.paragraph_format.space_before = TITLE_CONTACT_SPACE
        contact.paragraph_format.tab_stops.add_tab_stop(
            Inches(3.535), alignment=WD_TAB_ALIGNMENT.CENTER  # half of content width
        )
        contact.paragraph_format.tab_stops.add_tab_stop(
            LINE_WIDTH, alignment=WD_TAB_ALIGNMENT.RIGHT
        )
        contact.add_run(email + "\t" + phone + "\t" + location)
        for run in contact.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(CONTENT_FONT_SIZE)

    doc.add_paragraph()  # spacing

    # Section headers we expect
    section_headers = (
        "PROFILE",
        "WORK EXPERIENCE",
        "EDUCATION",
        "CERTIFICATIONS",
    )
    first_section = True
    in_certifications = False
    in_profile = False
    first_skill_after_summary = False
    in_work_experience = False
    first_job_in_work = True
    first_bullet_in_job = True  # first bullet per company → justify; rest → left
    CERT_LEFT_INDENT = Inches(0.25)
    section_heading_paragraphs = []  # for space_after (not zeroed at end)

    i = 4  # skip name, subtitle, blank, contact
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Section heading (all caps or known header)
        if stripped.upper() in section_headers or (
            stripped.isupper() and len(stripped) > 2 and stripped.isalpha()
        ):
            p = doc.add_paragraph()
            if not first_section:
                p.paragraph_format.space_before = SECTION_SPACE_BEFORE
            first_section = False
            in_certifications = False
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = FONT_NAME
            run.font.color.rgb = ACCENT_COLOR
            add_border_below(p)
            section_heading_paragraphs.append(p)
            in_certifications = stripped.upper() == "CERTIFICATIONS"
            in_profile = stripped.upper() == "PROFILE"
            if stripped.upper() == "WORK EXPERIENCE":
                in_work_experience = True
                first_job_in_work = True
            else:
                in_work_experience = False
            if in_profile:
                first_skill_after_summary = True
            i += 1
            continue

        # Certification line (cert name left, issuer right)
        if in_certifications and stripped:
            cert_name, issuer = split_cert_line(line)
            if cert_name or issuer:
                cert_para = doc.add_paragraph()
                cert_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                cert_para.paragraph_format.tab_stops.add_tab_stop(
                    LINE_WIDTH, alignment=WD_TAB_ALIGNMENT.RIGHT
                )
                cert_para.add_run(cert_name + "\t" + issuer)
                for run in cert_para.runs:
                    run.font.name = FONT_NAME
                i += 1
                continue

        # Standalone date/location line in work experience (e.g. last company with no job title in context)
        if in_work_experience and ("–" in stripped or " - " in stripped):
            if any(x in stripped for x in ["Present", "202", "201", "200", "Remote", "FL", "CA", "USA", "Gainesville", "Santa Clara"]):
                period, location = split_date_location(stripped)
                if period and location and not stripped.startswith("-"):
                    first_bullet_in_job = True
                    date_para = doc.add_paragraph()
                    if not first_job_in_work:
                        date_para.paragraph_format.space_before = COMPANY_BLOCK_SPACE
                    first_job_in_work = False
                    date_para.paragraph_format.tab_stops.add_tab_stop(
                        LINE_WIDTH, alignment=WD_TAB_ALIGNMENT.RIGHT
                    )
                    date_para.add_run(period + "\t" + location)
                    for run in date_para.runs:
                        run.font.name = FONT_NAME
                    i += 1
                    if i < len(lines) and lines[i].strip() and not lines[i].strip().startswith("-"):
                        desc_para = doc.add_paragraph(lines[i].strip())
                        desc_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        for run in desc_para.runs:
                            run.font.name = FONT_NAME
                        i += 1
                    continue

        # Job title line (e.g. "Senior Software Engineer – Wipro (Apple)")
        next_line = lines[i + 1].strip() if i + 1 < len(lines) else ""
        if next_line and (
            ("–" in stripped or " - " in stripped)
            and any(x in next_line for x in ["Present", "202", "Remote", "FL", "CA", "USA"])
        ):
            if in_work_experience:
                first_bullet_in_job = True  # new company: next bullet is first
            p = doc.add_paragraph()
            if in_work_experience and not first_job_in_work:
                p.paragraph_format.space_before = COMPANY_BLOCK_SPACE
            if in_work_experience:
                first_job_in_work = False
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(CONTENT_FONT_SIZE)
            run.font.name = FONT_NAME
            run.font.color.rgb = ACCENT_COLOR
            i += 1
            # Date/location line: period left, location right (one line with tab)
            if i < len(lines):
                period, location = split_date_location(lines[i])
                date_para = doc.add_paragraph()
                date_para.paragraph_format.space_before = Pt(0)
                date_para.paragraph_format.tab_stops.add_tab_stop(
                    LINE_WIDTH, alignment=WD_TAB_ALIGNMENT.RIGHT
                )
                date_para.add_run(period + "\t" + location)
                for run in date_para.runs:
                    run.font.name = FONT_NAME
                i += 1
            # Description paragraph (justified)
            if i < len(lines) and lines[i].strip() and not lines[i].strip().startswith("-"):
                desc_para = doc.add_paragraph(lines[i].strip())
                desc_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in desc_para.runs:
                    run.font.name = FONT_NAME
                i += 1
            continue

        # Bullet line
        if stripped.startswith("-") or stripped.startswith("•"):
            bullet_text = stripped.lstrip("- \t")
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(bullet_text)
            r.font.name = FONT_NAME
            if in_work_experience:
                r.font.size = Pt(9)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Inches(0.25)
            i += 1
            continue

        # Skill-style line: "Label: content"
        if ": " in stripped and not stripped.upper() in section_headers:
            idx = stripped.index(": ")
            label, content = stripped[: idx + 1], stripped[idx + 2 :]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = SKILL_HANGING_INDENT
            p.paragraph_format.first_line_indent = -SKILL_HANGING_INDENT
            if in_profile and first_skill_after_summary:
                p.paragraph_format.space_before = PROFILE_SKILLS_SPACE
                first_skill_after_summary = False
            r1 = p.add_run(label)
            r1.bold = True
            r1.font.name = FONT_NAME
            r2 = p.add_run(" " + content)
            r2.font.name = FONT_NAME
            i += 1
            continue

        # Regular paragraph (e.g. profile summary, education description)
        para = doc.add_paragraph(stripped)
        if in_profile:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in para.runs:
            run.font.name = FONT_NAME
        i += 1

    # No space after any paragraph; keep space after section headings
    for para in doc.paragraphs:
        para.paragraph_format.space_after = Pt(0)
    for p in section_heading_paragraphs:
        p.paragraph_format.space_after = HEADING_SPACE_AFTER

    # A4 paper and margins (top/bottom 0.75", left/right 0.6")
    a4_width = Inches(8.27)
    a4_height = Inches(11.69)
    for section in doc.sections:
        section.page_width = a4_width
        section.page_height = a4_height
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    doc.save(output_path)
    print(f"Saved: {output_path}")


def main():
    base = Path(__file__).resolve().parent
    context_path = base / "data/1/context.txt"
    output_path = base / "data/1/resume.docx"

    if not context_path.exists():
        print(f"Error: {context_path} not found.")
        return

    build_resume_docx(str(context_path), str(output_path))


if __name__ == "__main__":
    main()
