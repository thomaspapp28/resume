"""Template 3: Calibri-style, Professional Summary, Work Experience."""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import parse_xml

from app.core.config import BASE_DIR

ADDRESS_ICON_PATH = BASE_DIR / "assets" / "address-icon.png"

LINE_WIDTH = Inches(8.5 - 0.5 - 0.5)
FONT = "Calibri"
NAME_SIZE = Pt(20)
SUBTITLE_SIZE = Pt(12)
CONTENT_SIZE = Pt(10)
SECTION_SIZE = Pt(12)
ACCENT = RGBColor(0x1a, 0x86, 0xad)  # Professional dark navy

# PROFILE -> Professional Summary, WORK EXPERIENCE -> Work Experience


def build(context_path: str, output_path: str) -> None:
    with open(context_path, "r", encoding="utf-8") as f:
        text = f.read()
    lines = text.split("\n")

    def split_date_loc(line):
        parts = re.split(r"[\t ]{2,}", line.strip(), maxsplit=1)
        return (parts[0].strip() if parts else "", parts[1].strip() if len(parts) > 1 else "")

    def split_contact(line):
        parts = re.split(r"[\t ]{2,}", line.strip())
        parts = [p.strip() for p in parts if p.strip()]
        return (parts[0] if parts else "", parts[1] if len(parts) > 1 else "", parts[2] if len(parts) > 2 else "")

    def split_cert(line):
        content = line.strip().lstrip("-• \t")
        parts = re.split(r"\t|[\t ]{2,}", content, maxsplit=1)
        return (parts[0].strip() if parts else "", parts[1].strip() if len(parts) > 1 else "")

    def add_border_below(p):
        pPr = p._element.get_or_add_pPr()
        bdr = '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="12" w:space="1" w:color="000000"/></w:pBdr>'
        pPr.append(parse_xml(bdr))

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = CONTENT_SIZE
    style.paragraph_format.space_after = Pt(0)

    if lines:
        p = doc.add_paragraph()
        r = p.add_run(lines[0].strip())
        r.bold = True
        r.font.size = NAME_SIZE
        r.font.name = FONT
        r.font.color.rgb = ACCENT
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if len(lines) > 1:
        sub = doc.add_paragraph(lines[1].strip())
        sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in sub.runs:
            run.font.name = FONT
            run.font.size = SUBTITLE_SIZE
            run.font.color.rgb = ACCENT

    if len(lines) > 3 and lines[3].strip():
        email, phone, loc = split_contact(lines[3])
        c = doc.add_paragraph()
        c.paragraph_format.space_before = Pt(6)
        c.paragraph_format.tab_stops.add_tab_stop(Inches(3.535), alignment=WD_TAB_ALIGNMENT.CENTER)
        c.paragraph_format.tab_stops.add_tab_stop(LINE_WIDTH, alignment=WD_TAB_ALIGNMENT.RIGHT)
        r1 = c.add_run(f"\u2709 {email}\t\u260E {phone}\t")
        r1.font.name = FONT
        r1.font.size = CONTENT_SIZE
        r1.font.color.rgb = RGBColor(0, 0, 0)
        if ADDRESS_ICON_PATH.exists():
            r2 = c.add_run()
            r2.add_picture(str(ADDRESS_ICON_PATH), width=Pt(9), height=Pt(9))
            r3 = c.add_run(f" {loc}")
        else:
            r3 = c.add_run(f"\u2316 {loc}")
        r3.font.name = FONT
        r3.font.size = CONTENT_SIZE
        r3.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph()

    headers = ("PROFILE", "PROFESSIONAL SUMMARY", "WORK EXPERIENCE", "EDUCATION", "CERTIFICATIONS")
    first_section = True
    in_cert = in_profile = False
    first_skill = True
    in_work = False
    first_job = True
    section_paras = []
    i = 4

    while i < len(lines):
        line = lines[i]
        s = line.strip()
        if not s:
            i += 1
            continue

        su = s.upper()
        is_header = su in headers or (s.isupper() and len(s) > 2 and s.isalpha())
        if is_header:
            display = "Professional Summary" if su in ("PROFILE", "PROFESSIONAL SUMMARY") else (
                "Work Experience" if su == "WORK EXPERIENCE" else (
                "Education" if su == "EDUCATION" else "Certifications" if su == "CERTIFICATIONS" else s))
            p = doc.add_paragraph()
            if not first_section:
                p.paragraph_format.space_before = Pt(12)
            first_section = False
            in_cert = False
            r = p.add_run(display)
            r.bold = True
            r.font.size = SECTION_SIZE
            r.font.name = FONT
            r.font.color.rgb = ACCENT
            add_border_below(p)
            section_paras.append(p)
            in_cert = su == "CERTIFICATIONS"
            in_profile = su in ("PROFILE", "PROFESSIONAL SUMMARY")
            in_work = su == "WORK EXPERIENCE"
            if in_work:
                first_job = True
            if in_profile:
                first_skill = True
            i += 1
            continue

        if in_cert and s:
            name, issuer = split_cert(line)
            if name or issuer:
                cp = doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                cp.paragraph_format.tab_stops.add_tab_stop(LINE_WIDTH, alignment=WD_TAB_ALIGNMENT.RIGHT)
                cp.add_run(f"{name}\t{issuer}")
                for run in cp.runs:
                    run.font.name = FONT
                i += 1
                continue

        if in_work and ("–" in s or " - " in s):
            if any(x in s for x in ["Present", "202", "201", "200", "Remote", "FL", "CA", "USA", "Gainesville", "Santa Clara"]):
                period, loc = split_date_loc(s)
                if period and loc and not s.startswith("-"):
                    dp = doc.add_paragraph()
                    if not first_job:
                        dp.paragraph_format.space_before = Pt(10)
                    first_job = False
                    dp.paragraph_format.tab_stops.add_tab_stop(LINE_WIDTH, alignment=WD_TAB_ALIGNMENT.RIGHT)
                    dp.add_run(f"{period}\t{loc}")
                    for run in dp.runs:
                        run.font.name = FONT
                    i += 1
                    if i < len(lines) and lines[i].strip() and not lines[i].strip().startswith("-"):
                        d2 = doc.add_paragraph(lines[i].strip())
                        d2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        for run in d2.runs:
                            run.font.name = FONT
                        i += 1
                    continue

        next_s = lines[i + 1].strip() if i + 1 < len(lines) else ""
        if next_s and ("–" in s or " - " in s) and any(x in next_s for x in ["Present", "202", "Remote", "FL", "CA", "USA"]):
            p = doc.add_paragraph()
            if in_work and not first_job:
                p.paragraph_format.space_before = Pt(10)
            if in_work:
                first_job = False
            r = p.add_run(s)
            r.bold = True
            r.font.size = CONTENT_SIZE
            r.font.name = FONT
            r.font.color.rgb = ACCENT
            i += 1
            if i < len(lines):
                period, loc = split_date_loc(lines[i])
                dp = doc.add_paragraph()
                dp.paragraph_format.space_before = Pt(0)
                dp.paragraph_format.tab_stops.add_tab_stop(LINE_WIDTH, alignment=WD_TAB_ALIGNMENT.RIGHT)
                dp.add_run(f"{period}\t{loc}")
                for run in dp.runs:
                    run.font.name = FONT
                i += 1
            if i < len(lines) and lines[i].strip() and not lines[i].strip().startswith("-"):
                d2 = doc.add_paragraph(lines[i].strip())
                d2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in d2.runs:
                    run.font.name = FONT
                i += 1
            continue

        if s.startswith("-") or s.startswith("•"):
            bt = s.lstrip("- \t")
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(bt)
            r.font.name = FONT
            r.font.size = CONTENT_SIZE
            if in_work:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Inches(0.25)
            i += 1
            continue

        if ": " in s and su not in headers:
            idx = s.index(": ")
            label, content = s[: idx + 1], s[idx + 2 :]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = -Inches(0.5)
            if in_profile and first_skill:
                p.paragraph_format.space_before = Pt(8)
                first_skill = False
            r1 = p.add_run(label)
            r1.bold = True
            r1.font.name = FONT
            r2 = p.add_run(" " + content)
            r2.font.name = FONT
            i += 1
            continue

        p = doc.add_paragraph(s)
        if in_profile:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = FONT
            run.font.size = CONTENT_SIZE
        i += 1

    for para in doc.paragraphs:
        para.paragraph_format.space_after = Pt(0)
    for p in section_paras:
        p.paragraph_format.space_after = Pt(6)

    for sec in doc.sections:
        sec.page_width = Inches(8.5)
        sec.page_height = Inches(12)
        sec.top_margin = Inches(0.5)
        sec.bottom_margin = Inches(0.5)
        sec.left_margin = Inches(0.5)
        sec.right_margin = Inches(0.5)

    doc.save(output_path)


CONFIG = {"id": 3, "name": "template3"}
