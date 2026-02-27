"""Template 5: Roboto name, centered section headers, A4 page."""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import parse_xml

LINE_WIDTH = Inches(8.27 - 0.5 - 0.5)
ICON_MAIL = "\u2709"
ICON_PHONE = "\u260E"
ICON_ADDRESS = "\u2302"
SYMBOL_FONT = "Segoe UI Symbol"
FONT = "Times New Roman"
NAME_FONT = "Roboto"
NAME_SIZE = Pt(24)
SUBTITLE_SIZE = Pt(12)
CONTENT_SIZE = Pt(10)
BULLET_SIZE = Pt(9.5)
SECTION_SIZE = Pt(10)
ACCENT = RGBColor(0x1e, 0x3a, 0x5f)  # Professional dark navy


def build(context_path: str, output_path: str) -> None:
    with open(context_path, "r", encoding="utf-8") as f:
        text = f.read()
    lines = text.split("\n")

    def split_date_loc(line):
        parts = re.split(r"[\t ]{2,}", line.strip(), maxsplit=1)
        return (parts[0].strip() if parts else "", parts[1].strip() if len(parts) > 1 else "")

    def split_contact(line):
        parts = [p.strip() for p in line.strip().split("\t")]
        return (parts[0] if len(parts) > 0 else "", parts[1] if len(parts) > 1 else "", parts[2] if len(parts) > 2 else "")

    def split_cert(line):
        content = line.strip().lstrip("-• \t")
        parts = re.split(r"\t|[\t ]{2,}", content, maxsplit=1)
        return (parts[0].strip() if parts else "", parts[1].strip() if len(parts) > 1 else "")

    def add_border_below(p):
        pPr = p._element.get_or_add_pPr()
        bdr = '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="12" w:space="1" w:color="000000"/></w:pBdr>'
        pPr.append(parse_xml(bdr))

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = CONTENT_SIZE
    style.paragraph_format.space_after = Pt(0)

    if lines:
        p = doc.add_paragraph()
        r = p.add_run(lines[0].strip())
        r.bold = True
        r.font.size = NAME_SIZE
        try:
            r.font.name = NAME_FONT
        except Exception:
            r.font.name = "Calibri"
        r.font.color.rgb = ACCENT
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if len(lines) > 1:
        sub = doc.add_paragraph()
        r = sub.add_run(lines[1].strip())
        r.bold = True
        r.font.size = SUBTITLE_SIZE
        try:
            r.font.name = NAME_FONT
        except Exception:
            r.font.name = "Calibri"
        r.font.color.rgb = ACCENT
        sub.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if len(lines) > 3 and lines[3].strip():
        email, phone, loc = split_contact(lines[3])
        c = doc.add_paragraph()
        c.paragraph_format.space_before = Pt(6)
        c.paragraph_format.tab_stops.add_tab_stop(Inches(3.535), alignment=WD_TAB_ALIGNMENT.CENTER)
        c.paragraph_format.tab_stops.add_tab_stop(LINE_WIDTH, alignment=WD_TAB_ALIGNMENT.RIGHT)
        black = RGBColor(0, 0, 0)
        for icon, text, sep in [(ICON_MAIL, email, "\t"), (ICON_PHONE, phone, "\t"), (ICON_ADDRESS, loc, "")]:
            ri = c.add_run(f"{icon} ")
            ri.font.name = SYMBOL_FONT
            ri.font.size = CONTENT_SIZE
            ri.font.color.rgb = black
            rt = c.add_run(f"{text}{sep}")
            rt.font.name = FONT
            rt.font.size = CONTENT_SIZE
            rt.font.color.rgb = black

    doc.add_paragraph()

    headers = ("PROFILE", "WORK EXPERIENCE", "EXPERIENCE", "PROFESSIONAL EXPERIENCE", "EDUCATION", "CERTIFICATIONS")
    first_section = True
    in_cert = in_profile = False
    first_skill = True
    in_work = False
    first_job = True
    section_paras = []
    i = 4

    while i < len(lines):
        line = lines[i]
        s = line.strip()
        if not s:
            i += 1
            continue

        if s.upper() in headers or (s.isupper() and len(s) > 2 and s.isalpha()):
            p = doc.add_paragraph()
            if not first_section:
                p.paragraph_format.space_before = Pt(16)
            first_section = False
            in_cert = False
            r = p.add_run(s.upper())
            r.bold = True
            r.font.size = SECTION_SIZE
            r.font.name = FONT
            r.font.color.rgb = ACCENT
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_border_below(p)
            section_paras.append(p)
            in_cert = s.upper() == "CERTIFICATIONS"
            in_profile = s.upper() == "PROFILE"
            in_work = s.upper() in ("WORK EXPERIENCE", "EXPERIENCE", "PROFESSIONAL EXPERIENCE")
            if in_work:
                first_job = True
            if in_profile:
                first_skill = True
            i += 1
            continue

        if in_cert and s:
            name, issuer = split_cert(line)
            if name or issuer:
                cp = doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                cp.paragraph_format.tab_stops.add_tab_stop(LINE_WIDTH, alignment=WD_TAB_ALIGNMENT.RIGHT)
                cp.add_run(f"{name}\t{issuer}")
                for run in cp.runs:
                    run.font.name = FONT
                i += 1
                continue

        if in_work and ("–" in s or " - " in s):
            if any(x in s for x in ["Present", "202", "201", "200", "Remote", "FL", "CA", "USA", "Gainesville", "Santa Clara"]):
                period, loc = split_date_loc(s)
                if period and loc and not s.startswith("-"):
                    dp = doc.add_paragraph()
                    if not first_job:
                        dp.paragraph_format.space_before = Pt(14)
                    first_job = False
                    dp.paragraph_format.tab_stops.add_tab_stop(LINE_WIDTH, alignment=WD_TAB_ALIGNMENT.RIGHT)
                    dp.add_run(f"{period}\t{loc}")
                    for run in dp.runs:
                        run.font.name = FONT
                    i += 1
                    if i < len(lines) and lines[i].strip() and not lines[i].strip().startswith("-"):
                        d2 = doc.add_paragraph(lines[i].strip())
                        d2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        for run in d2.runs:
                            run.font.name = FONT
                        i += 1
                    continue

        next_s = lines[i + 1].strip() if i + 1 < len(lines) else ""
        if next_s and ("–" in s or " - " in s) and any(x in next_s for x in ["Present", "202", "Remote", "FL", "CA", "USA"]):
            p = doc.add_paragraph()
            if in_work and not first_job:
                p.paragraph_format.space_before = Pt(14)
            if in_work:
                first_job = False
            r = p.add_run(s)
            r.bold = True
            r.font.size = CONTENT_SIZE
            try:
                r.font.name = NAME_FONT
            except Exception:
                r.font.name = "Calibri"
            r.font.color.rgb = ACCENT
            i += 1
            if i < len(lines):
                period, loc = split_date_loc(lines[i])
                dp = doc.add_paragraph()
                dp.paragraph_format.space_before = Pt(0)
                dp.paragraph_format.tab_stops.add_tab_stop(LINE_WIDTH, alignment=WD_TAB_ALIGNMENT.RIGHT)
                dp.add_run(f"{period}\t{loc}")
                for run in dp.runs:
                    run.font.name = FONT
                i += 1
            if i < len(lines) and lines[i].strip() and not lines[i].strip().startswith("-"):
                d2 = doc.add_paragraph(lines[i].strip())
                d2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in d2.runs:
                    run.font.name = FONT
                i += 1
            continue

        if s.startswith("-") or s.startswith("•"):
            bt = s.lstrip("- \t")
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(bt)
            r.font.name = FONT
            r.font.size = BULLET_SIZE
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Inches(0.25)
            i += 1
            continue

        if ": " in s and s.upper() not in headers:
            idx = s.index(": ")
            label, content = s[: idx + 1], s[idx + 2 :]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = -Inches(0.5)
            if in_profile and first_skill:
                p.paragraph_format.space_before = Pt(8)
                first_skill = False
            r1 = p.add_run(label)
            r1.bold = True
            r1.font.name = FONT
            r2 = p.add_run(" " + content)
            r2.font.name = FONT
            i += 1
            continue

        p = doc.add_paragraph(s)
        if in_profile:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = FONT
        i += 1

    for para in doc.paragraphs:
        para.paragraph_format.space_after = Pt(0)
    for p in section_paras:
        p.paragraph_format.space_after = Pt(10)

    for sec in doc.sections:
        sec.page_width = Inches(8.27)
        sec.page_height = Inches(11.69)
        sec.top_margin = Inches(0.875)
        sec.bottom_margin = Inches(0.625)
        sec.left_margin = Inches(0.5)
        sec.right_margin = Inches(0.5)

    doc.save(output_path)


CONFIG = {"id": 5, "name": "template5"}
